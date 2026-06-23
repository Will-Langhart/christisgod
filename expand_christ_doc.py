from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


SRC = Path("/Users/lang/Downloads/Christ_Is_God_The_Divinity_of_Christ.docx")
OUT = Path("/Users/lang/christisgod/Christ_Is_God_The_Divinity_of_Christ_expanded.docx")


def para_text(p):
    return p.text.strip()


def find_para(doc, startswith):
    for p in doc.paragraphs:
        if para_text(p).startswith(startswith):
            return p
    raise ValueError(f"paragraph not found: {startswith}")


def find_heading(doc, startswith, level=None):
    expected = f"Heading {level}" if level else "Heading"
    for p in doc.paragraphs:
        style_name = p.style.name if p.style is not None else ""
        if expected in style_name and para_text(p).startswith(startswith):
            return p
    raise ValueError(f"heading not found: {startswith}")


def clear_para(p, text):
    p.clear()
    p.add_run(text)


def set_keep_with_next(p):
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn("w:keepNext")) is None:
        pPr.append(OxmlElement("w:keepNext"))


def insert_before(ref_para, text="", style=None):
    new_p = OxmlElement("w:p")
    ref_para._p.addprevious(new_p)
    p = ref_para.__class__(new_p, ref_para._parent)
    if style:
        p.style = STYLE_BY_NAME.get(style, style)
    if text:
        p.add_run(text)
    if style and "Heading" in style:
        set_keep_with_next(p)
    return p


def insert_after(ref_para, text="", style=None):
    new_p = OxmlElement("w:p")
    ref_para._p.addnext(new_p)
    p = ref_para.__class__(new_p, ref_para._parent)
    if style:
        p.style = STYLE_BY_NAME.get(style, style)
    if text:
        p.add_run(text)
    if style and "Heading" in style:
        set_keep_with_next(p)
    return p


def add_block_before(ref_para, items):
    current = ref_para
    made = []
    for text, style in reversed(items):
        p = insert_before(current, text, style)
        made.append(p)
        current = p
    return list(reversed(made))


def copy_table_before(ref_para, table):
    ref_para._p.addprevious(deepcopy(table._tbl))


def roman_contents():
    return [
        "I. Introduction — Why the Divinity of Christ Matters",
        "II. Romans 9:5 — The Anchor Text",
        "III. “Jesus Is God” — The Direct Witness of Scripture",
        "IV. Old Testament Theophanies — The Angel of the LORD",
        "V. The Divine Works of Christ",
        "VI. The Indirect but Unmistakable Claim",
        "VII. Worship Given to Christ",
        "VIII. The Omnipresence of Christ",
        "IX. The Holy Trinity — A Comparative Table",
        "X. The Witness of the Early Church Fathers",
        "XI. Islam's View of Christ",
        "XII. Anticipating Objections",
        "XIII. The Patristic and Liturgical Witness of the Coptic Church",
        "XIV. Conclusion — From Confession to Communion",
        "Sources Consulted",
    ]


doc = Document(SRC)
STYLE_BY_NAME = {}
for existing in doc.paragraphs:
    if existing.style is not None and existing.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
        STYLE_BY_NAME.setdefault(existing.style.name, existing.style)

# Page setup: keep the document's existing feel, but ensure tables have workable width.
for section in doc.sections:
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

# Public-facing generalization and small corrections.
for p in doc.paragraphs:
    txt = para_text(p)
    if "Lang's experience" in txt:
        txt = txt.replace("Lang's experience", "many contemporary apologetic conversations")
    if "Lang's debates with Jehovah's Witnesses and other non-trinitarian interlocutors" in txt:
        txt = txt.replace(
            "Lang's debates with Jehovah's Witnesses and other non-trinitarian interlocutors",
            "conversations with Jehovah's Witnesses, Muslims, unitarians, and other non-trinitarian interlocutors",
        )
    if txt != para_text(p):
        clear_para(p, txt)
    if "The accompanying catechism slide" in txt:
        clear_para(
            p,
            txt.replace(
                "The accompanying catechism slide (reproduced and corrected below in table form)",
                "The comparative table below",
            ),
        )
    if txt.startswith("Cyril's twenty-fourth Pope"):
        clear_para(
            p,
            txt.replace(
                "Cyril's twenty-fourth Pope of Alexandria's defense",
                "As the twenty-fourth Pope of Alexandria, Cyril's defense",
            ),
        )

# Rename existing numbered sections after new insertions.
renames = {
    "IV. The Divine Works of Christ": "V. The Divine Works of Christ",
    "V. The Indirect but Unmistakable Claim": "VI. The Indirect but Unmistakable Claim",
    "VI. The Omnipresence of Christ": "VIII. The Omnipresence of Christ",
    "VII. The Holy Trinity — A Comparative Table": "IX. The Holy Trinity — A Comparative Table",
    "VIII. The Witness of the Early Church Fathers": "X. The Witness of the Early Church Fathers",
    "IX. Anticipating Objections": "XII. Anticipating Objections",
    "X. The Patristic and Liturgical Witness of the Coptic Church": "XIII. The Patristic and Liturgical Witness of the Coptic Church",
    "XI. Conclusion — From Confession to Communion": "XIV. Conclusion — From Confession to Communion",
}
for p in doc.paragraphs:
    t = para_text(p)
    if t in renames:
        clear_para(p, renames[t])

# Rewrite contents list.
contents = find_para(doc, "Contents")
cursor = contents
for _ in range(11):
    nxt = cursor._p.getnext()
    if nxt is not None:
        cursor._p.getparent().remove(nxt)
for line in roman_contents():
    cursor = insert_after(cursor, line)

# Deepen Romans 9:5 and textual-critical discussion.
romans_anchor = find_para(doc, "Some modern translations punctuate")
insert_after(
    romans_anchor,
    "Two observations strengthen the traditional reading. First, the phrase “according to the flesh” naturally invites a second, higher statement about who Christ is beyond His Davidic and Israelite lineage. Second, Paul elsewhere reserves doxological language for God, but here the grammar most naturally binds “who is over all” and “God blessed forever” to Christ Himself. Tertullian, writing against Praxeas, explicitly appeals to Romans 9:5 and says that when Christ is considered in Himself, Paul can call Him God. The verse is therefore not a late medieval polemical invention but an early Christian proof-text already active in second- and third-century trinitarian argument.",
)

timothy = find_para(doc, "Paul (or the hymn he is quoting) confesses")
clear_para(
    timothy,
    "Paul (or the hymn he is quoting) confesses: “Great is the mystery of godliness: God was manifested in the flesh, justified in the Spirit, seen of angels, preached unto the Gentiles, believed on in the world, received up into glory.” A textual note belongs here. The King James reading “God was manifested in the flesh” reflects the later Byzantine tradition and the theological confession is entirely orthodox; nevertheless, many modern critical editions read “He who” or “who.” Even on that reading, the hymn still moves from incarnation to angelic witness, worldwide proclamation, faith among the nations, and reception into glory. The doctrine does not depend on one disputed pronoun, because the wider New Testament already identifies the incarnate One as the divine Son, the Creator, the Judge, and the Lord who receives worship.",
)

burgon = find_para(doc, "In fairness to the historical record")
insert_after(
    burgon,
    "This means the Burgon-style table is best used as cumulative corroboration rather than as a stand-alone proof. It shows that the New Testament was deeply embedded in the preaching, controversy, commentary, and liturgy of the early Church; it should not be pressed as though the Fathers alone reproduce every verse in a mechanically complete critical edition. Its real force is historical atmosphere: the same Church that preserved, quoted, preached, and suffered for these writings also confessed Christ as God.",
)

# Insert Old Testament theophanies before Divine Works.
divine_works = find_heading(doc, "V. The Divine Works of Christ", 1)
add_block_before(
    divine_works,
    [
        ("IV. Old Testament Theophanies — The Angel of the LORD", "Heading 1"),
        (
            "The divinity of Christ does not appear suddenly in the New Testament as though the apostles invented a new doctrine. The Old Testament already contains a mysterious pattern: the Angel of the LORD appears as a messenger distinct from God, and yet speaks as God, receives divine fear, reveals the divine Name, and is identified by the narrator with the LORD Himself. Christian tradition has often seen in these appearances not the created angelic hosts, but pre-incarnate manifestations of the Word who would later become flesh.",
            None,
        ),
        ("Genesis 16; Genesis 22; Exodus 3 — Distinct, Yet Divine", "Heading 2"),
        (
            "In Genesis 16 the Angel of the LORD speaks to Hagar in the first person as the One who will multiply her descendants, and Hagar responds by naming the LORD who spoke to her. In Genesis 22, the Angel calls from heaven and says, “By Myself I have sworn,” language proper to God alone. In Exodus 3, the Angel of the LORD appears in the burning bush, yet the text immediately says that God called to Moses from the bush; Moses hides his face because he is afraid to look upon God. The messenger is not a rival deity and not a mere creature speaking independently, but the divine Presence personally revealed.",
            None,
        ),
        ("Judges 13 — Manoah's Fear", "Heading 2"),
        (
            "The parents of Samson encounter the Angel of the LORD, who accepts an offering ascending in the flame of the altar. Manoah then says, “We shall surely die, because we have seen God.” His wife corrects the conclusion that they will die, but not the identification that the encounter was divine. This pattern matters for Christology because John later writes that no one has seen God the Father in His unveiled essence, while the only-begotten Son makes Him known. The Son is therefore the visible self-revelation of the invisible God.",
            None,
        ),
    ],
)

# Insert worship section before Omnipresence.
omnipresence = find_heading(doc, "VIII. The Omnipresence of Christ", 1)
add_block_before(
    omnipresence,
    [
        ("VII. Worship Given to Christ", "Heading 1"),
        (
            "Scripture does not merely call Christ by divine titles or attribute divine works to Him; it shows Him receiving the kind of reverence that belongs to God. This is especially important because both the Old and New Testaments fiercely reject worship offered to creatures. Angels refuse it. Apostles refuse it. God alone is to be served and adored. Yet Jesus receives worship without rebuke, and the heavenly liturgy centers on the Lamb.",
            None,
        ),
        ("Matthew 14:33; Matthew 28:9, 17 — Worship After Revelation", "Heading 2"),
        (
            "After Jesus walks on the sea and stills the fear of the disciples, those in the boat worship Him and confess, “Truly You are the Son of God.” After the Resurrection, the women take hold of His feet and worship Him; the eleven also worship on the mountain in Galilee. These scenes are not accidental gestures of respect. They follow revelations of divine authority over creation and death.",
            None,
        ),
        ("John 9:35–38 and Hebrews 1:6 — Worship Commanded", "Heading 2"),
        (
            "When the man born blind confesses faith in the Son of God, he worships Jesus, and Jesus accepts it. Hebrews goes further: when the Firstborn is brought into the world, the angels of God are commanded to worship Him. A created angel may not receive worship, but the Son is worshiped by angels because He is not a fellow creature. He is the radiance of the Father's glory and the exact image of His hypostasis.",
            None,
        ),
        ("Revelation 5 — The Lamb on the Throne", "Heading 2"),
        (
            "Revelation 5 places the slain and risen Lamb inside the worship of heaven. Blessing, honor, glory, and power are given “to Him who sits on the throne, and to the Lamb.” The doxology does not split worship between God and a creature; it reveals the one divine worship shared by the Father and the Son. This is why Christian prayer, hymnody, and liturgy are christological without becoming idolatrous.",
            None,
        ),
    ],
)

# Insert Islam section before objections.
objections = find_heading(doc, "XII. Anticipating Objections", 1)
add_block_before(
    objections,
    [
        ("XI. Islam's View of Christ", "Heading 1"),
        (
            "Because this study is intended for public apologetic use, it is worth naming Islam's view of Christ with accuracy and charity. The Qur'an honors Jesus as Messiah, son of Mary, prophet, messenger, Word from God, and spirit from Him; it affirms His virginal conception and gives Mary an unusually honored place. At the same time, it explicitly denies that Jesus is God, denies divine sonship as Christianity understands it, rejects the Trinity, and presents Jesus as a servant and prophet who calls Israel to worship Allah alone.",
            None,
        ),
        ("The Qur'anic Affirmations", "Heading 2"),
        (
            "Qur'an 4:171 calls Jesus the Messiah, son of Mary, a messenger of Allah, His Word cast to Mary, and a spirit from Him. Qur'an 19:30 presents Jesus as saying that he is the servant of Allah and a prophet. These affirmations are significant for conversation: Islam does not treat Jesus as a false prophet or ordinary man. It preserves a high view of His birth, mission, and eschatological importance.",
            None,
        ),
        ("The Qur'anic Denials", "Heading 2"),
        (
            "The same passages mark the decisive disagreement. Qur'an 4:171 warns against saying “three” and denies that God has a son. Qur'an 5:72 rejects the claim that Allah is the Messiah, son of Mary. Qur'an 5:116 portrays Jesus denying that he told people to take him and his mother as deities besides Allah. A Christian response should therefore avoid caricature: the real dispute is not whether Jesus is honored, but whether the apostolic witness is true that the eternal Word became flesh and that the Son shares the Father's divine identity.",
            None,
        ),
        ("A Christian Point of Contact", "Heading 2"),
        (
            "The strongest bridge is the title “Word.” In Islam, the title is usually interpreted as God's creative command by which Jesus was miraculously conceived. In John, however, the Word is personal, eternal, with God, and God, and then becomes flesh. The conversation can therefore move from shared reverence for Jesus to the central question: is the Word merely a command that produces Jesus, or is the Word Himself the eternal Son who entered the womb of the Virgin Mary for our salvation?",
            None,
        ),
    ],
)

# Expand objections before closing paragraph of that section.
objection_end = find_para(doc, "None of this is offered")
add_block_before(
    objection_end,
    [
        ("“No one knows the day or hour, not even the Son” (Mark 13:32)", "Heading 3"),
        (
            "This text belongs to the same incarnational pattern as John 14:28. The eternal Son truly assumed a human mind, human will, and human condition; He did not merely wear humanity as a costume. Orthodox Christology therefore confesses both the divine omniscience of the Word and the real human consciousness of Jesus. The verse is not an argument that Christ is a creature, but a witness that the incarnation was genuine.",
            None,
        ),
        ("“Why do you call Me good? No one is good but God alone” (Mark 10:18)", "Heading 3"),
        (
            "Jesus is not denying His goodness; He is pressing the rich young ruler to understand the meaning of his own words. If goodness belongs ultimately to God, then calling Jesus good should lead the questioner upward, not backward. The passage functions as a spiritual diagnosis: the man uses reverent language but has not yet perceived who stands before him.",
            None,
        ),
        ("“The only true God, and Jesus Christ whom You have sent” (John 17:3)", "Heading 3"),
        (
            "In context, Jesus has just asked the Father to glorify Him with the glory He had with the Father before the world existed. John 17:3 distinguishes the Father and the Son personally; it does not exclude the Son from the divine identity. The same Gospel that records this prayer begins by saying that the Word was God and ends with Thomas confessing Jesus as “my Lord and my God.”",
            None,
        ),
        ("“My God and your God” (John 20:17)", "Heading 3"),
        (
            "The risen Christ speaks as the incarnate Son and true man, leading His brethren to the Father. The phrase does not erase His divine sonship; it reveals the grace of adoption. The Father is Christ's Father by eternal generation, and our Father by union with Christ. Likewise, the Father is His God according to His assumed humanity, while Christ remains the eternal Word according to His divinity.",
            None,
        ),
    ],
)

# Deepen Coptic/patristic section.
ignatius = find_para(doc, "Writing within a decade")
insert_after(
    ignatius,
    "In the Epistle to the Ephesians, Ignatius describes Christ as “God existing in flesh,” a compact phrase that already guards both sides of the mystery: the One who appears in flesh is truly God, and the flesh in which He appears is real. This is why Ignatius can oppose docetism and confess Christ's deity in the same breath.",
)
athanasius = find_para(doc, "As the unwavering champion")
insert_after(
    athanasius,
    "Athanasius' famous line, “He was made man that we might be made God,” is not rhetorical excess. It is the Alexandrian grammar of salvation: only God can unite humanity to God, and only a true human life can heal human nature from within.",
)
cyril = find_para(doc, "As the twenty-fourth Pope")
insert_after(
    cyril,
    "The title Theotokos is therefore not primarily a Marian embellishment but a Christological safeguard. It says that the one born of Mary is the same divine Person who is eternally begotten of the Father. Mary does not originate the Godhead; she bears the incarnate Word according to the flesh.",
)

# Add sources before final cross.
cross = None
for p in reversed(doc.paragraphs):
    if para_text(p) == "✠":
        cross = p
        break
if cross is None:
    cross = doc.add_paragraph()

add_block_before(
    cross,
    [
        ("Sources Consulted", "Heading 1"),
        (
            "Biblical texts: Romans 9:5; John 1:1–14; John 5; John 8; John 10; John 17; John 20; Hebrews 1; Revelation 5; Genesis 16; Genesis 22; Exodus 3; Judges 13; Isaiah 40; Isaiah 44; Isaiah 45.",
            None,
        ),
        (
            "Patristic sources checked online: St. Ignatius of Antioch, Epistle to the Ephesians (https://www.newadvent.org/fathers/0104.htm); Tertullian, Against Praxeas (https://www.newadvent.org/fathers/0317.htm); St. Athanasius, On the Incarnation of the Word (https://www.newadvent.org/fathers/2802.htm); conciliar and Cyrilline material surrounding Ephesus and the title Theotokos (https://www.newadvent.org/fathers/3810.htm).",
            None,
        ),
        (
            "Islam section checked against Qur'an 4:171 (https://quran.com/4/171); 5:72 (https://quran.com/5/72); 5:116 (https://quran.com/5/116); 19:30 (https://quran.com/19/30), using Quran.com for Arabic text and English translation.",
            None,
        ),
        (
            "Text-critical caution: the Burgon patristic quotation table is retained as a historical-apologetic aid, but the surrounding prose now warns against using the numbers as a precise or exhaustive reconstruction claim.",
            None,
        ),
    ],
)

# Keep the Trinity table from sprawling too far.
for table in doc.tables:
    table.autofit = True

doc.save(OUT)
print(OUT)
