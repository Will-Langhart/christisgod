from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

SRC = Path("/Users/lang/christisgod/Christ_Is_God_The_Divinity_of_Christ_expanded.docx")
OUT = Path("/Users/lang/christisgod/Christ_Is_God_The_Divinity_of_Christ_expanded_researched.docx")


def text(p):
    return p.text.strip()


def collect_styles(doc):
    styles = {}
    for p in doc.paragraphs:
        if p.style is not None and p.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            styles.setdefault(p.style.name, p.style)
    return styles


def keep_with_next(p):
    p_pr = p._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def find_heading(doc, prefix, level=None):
    expected = f"Heading {level}" if level else "Heading"
    for p in doc.paragraphs:
        style_name = p.style.name if p.style is not None else ""
        if expected in style_name and text(p).startswith(prefix):
            return p
    raise ValueError(f"heading not found: {prefix}")


def find_para(doc, prefix):
    for p in doc.paragraphs:
        if text(p).startswith(prefix):
            return p
    raise ValueError(f"paragraph not found: {prefix}")


def insert_before(ref, content="", style=None):
    node = OxmlElement("w:p")
    ref._p.addprevious(node)
    p = ref.__class__(node, ref._parent)
    if style:
        p.style = STYLES.get(style, style)
    if content:
        p.add_run(content)
    if style and "Heading" in style:
        keep_with_next(p)
    return p


def insert_after(ref, content="", style=None):
    node = OxmlElement("w:p")
    ref._p.addnext(node)
    p = ref.__class__(node, ref._parent)
    if style:
        p.style = STYLES.get(style, style)
    if content:
        p.add_run(content)
    if style and "Heading" in style:
        keep_with_next(p)
    return p


def block_after(anchor, items):
    cur = anchor
    for content, style in items:
        cur = insert_after(cur, content, style)
    return cur


def block_before(anchor, items):
    cur = anchor
    made = []
    for content, style in reversed(items):
        cur = insert_before(cur, content, style)
        made.append(cur)
    return list(reversed(made))


def clear(p, content):
    p.clear()
    p.add_run(content)


def rewrite_contents(doc):
    contents = find_heading(doc, "Contents", 1)
    cur = contents
    # Remove old manually-written contents lines until the first real section heading.
    while True:
        nxt = cur._p.getnext()
        if nxt is None:
            break
        candidate = cur.__class__(nxt, cur._parent)
        style_name = candidate.style.name if candidate.style is not None else ""
        if "Heading 1" in style_name and text(candidate).startswith("I. Introduction"):
            break
        cur._p.getparent().remove(nxt)
    lines = [
        "I. Introduction — Why the Divinity of Christ Matters",
        "II. Romans 9:5 — The Anchor Text",
        "III. “Jesus Is God” — The Direct Witness of Scripture",
        "IV. Old Testament Theophanies — The Angel of the LORD",
        "V. The Divine Works of Christ",
        "VI. The Indirect but Unmistakable Claim",
        "VII. Worship Given to Christ",
        "VIII. The Omnipresence of Christ",
        "IX. The Holy Trinity — A Comparative Table",
        "X. The Witness of the Early Church Fathers",
        "XI. Islam's View of Christ",
        "XII. Anticipating Objections",
        "XIII. The Patristic and Liturgical Witness of the Coptic Church",
        "XIV. A Compact Apologetic Case",
        "XV. Conclusion — From Confession to Communion",
        "Sources Consulted",
    ]
    cur = contents
    for line in lines:
        cur = insert_after(cur, line)


doc = Document(SRC)
STYLES = collect_styles(doc)

for section in doc.sections:
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

rewrite_contents(doc)

# Renumber the existing conclusion after adding a new section.
old_conclusion = find_heading(doc, "XIV. Conclusion", 1)
clear(old_conclusion, "XV. Conclusion — From Confession to Communion")

# More direct biblical evidence.
hebrews = find_heading(doc, "Hebrews 1:8", 2)
block_after(
    hebrews,
    [
        (
            "1 Corinthians 8:6 — Jesus Inside the Shema",
            "Heading 2",
        ),
        (
            "Paul's confession in 1 Corinthians 8:6 is one of the most compact and powerful arguments for Christ's deity. He begins from Jewish monotheism: “there is no other God but one.” Then, instead of placing Jesus outside that confession, he reshapes the language around “one God, the Father” and “one Lord Jesus Christ.” The Father is the source “of whom are all things”; the Son is the mediator “through whom are all things.” This is not a demotion of Jesus into the category of creature. It is a christological rereading of the Shema: the one God of Israel is confessed with the Father and the Lord Jesus Christ without abandoning monotheism.",
            None,
        ),
        (
            "Philippians 2:5–11 and Isaiah 45 — The Divine Name and Universal Worship",
            "Heading 2",
        ),
        (
            "Philippians 2 says that Christ existed in the form of God, humbled Himself in true humanity, and was then exalted so that every knee should bow and every tongue confess that Jesus Christ is Lord. Paul is deliberately echoing Isaiah 45, where YHWH alone declares that He is God and there is no other, and then swears that every knee will bow to Him. The apostolic claim is astonishing: the worship promised to YHWH is rendered to Jesus, and this glorifies the Father rather than competing with Him. The Son shares the divine Name and receives the divine homage because He belongs within the identity of the one God.",
            None,
        ),
        (
            "Colossians 1:15–20 — Creator, Sustainer, Reconciler",
            "Heading 2",
        ),
        (
            "Colossians calls Christ “the image of the invisible God” and “firstborn over all creation,” then immediately explains that all things were created by Him, through Him, and for Him. “Firstborn” therefore cannot mean “first creature,” because the Son is placed on the Creator side of the Creator-creature distinction. All visible and invisible powers are created through Him; He is before all things; in Him all things hold together. The same passage moves from creation to redemption: the One through whom all things were made is the One through whose blood all things are reconciled.",
            None,
        ),
    ],
)

# Add more works and devotional evidence.
maker_law = find_heading(doc, "6. Christ Is Maker of the Law", 2)
block_after(
    maker_law,
    [
        ("7. Christ Sustains All Things", "Heading 2"),
        (
            "Creation is not only an event in the past; it is also God's continual sustaining of all that exists. Hebrews 1 says that the Son upholds all things by the word of His power, and Colossians 1 says that in Him all things hold together. This is not the work of an angelic deputy who merely executes orders. To sustain being itself is a divine act. Every breath, atom, power, throne, and principality remains in existence because the Son is not one item within creation but the Lord through whom creation is held in being.",
            None,
        ),
        ("8. Christ Is Invoked in Prayer", "Heading 2"),
        (
            "The earliest Christians do not merely speak about Jesus; they call upon Him. Stephen prays, “Lord Jesus, receive my spirit” (Acts 7:59). Paul describes Christians as those who call upon the name of our Lord Jesus Christ (1 Corinthians 1:2). The Aramaic prayer Maranatha, “Our Lord, come,” preserved in 1 Corinthians 16:22 and echoed at the end of Revelation, shows devotion to Jesus embedded in the worshiping life of the apostolic Church. Prayer offered to Christ is not an ornamental detail; it is lived Christology.",
            None,
        ),
    ],
)

# Deepen worship and early corroboration.
revelation5 = find_heading(doc, "Revelation 5", 2)
block_after(
    revelation5,
    [
        ("Maranatha and the Earliest Prayer to Jesus", "Heading 2"),
        (
            "Maranatha matters because it is Aramaic, not later Greek theological vocabulary. It likely reaches back into the earliest Semitic-speaking Christian communities. A prayer asking the Lord Jesus to come assumes that He is alive, heavenly, able to answer, and worthy to be addressed in worship. This is not a fourth-century invention projected backward; it is embedded in Paul's first-century letters.",
            None,
        ),
        ("Pliny the Younger — A Pagan Witness to Christian Worship", "Heading 2"),
        (
            "Around AD 112, the Roman governor Pliny wrote to Trajan that Christians met before dawn and sang a hymn to Christ “as to a god.” Pliny was not trying to defend Christianity; he was interrogating Christians as a Roman official. That makes his testimony valuable: from outside the Church, he confirms that worship of Christ was already a recognizable marker of Christian identity at the beginning of the second century.",
            None,
        ),
        ("Justin Martyr — Worship of Father, Son, and Spirit", "Heading 2"),
        (
            "By the mid-second century, St. Justin Martyr could explain Christian worship to the Roman emperor by saying that Christians worship the true God, the Son who came from Him, and the prophetic Spirit. Justin's language is not yet the later technical vocabulary of Nicaea, but the devotional pattern is unmistakably trinitarian in seed: Christian worship is directed to the Father through and with the Son in the Holy Spirit.",
            None,
        ),
    ],
)

# More objections.
objection_end = find_para(doc, "None of this is offered")
block_before(
    objection_end,
    [
        ("“The beginning of the creation of God” (Revelation 3:14)", "Heading 3"),
        (
            "The Greek word arche can mean beginning, source, ruler, or originating principle. In Revelation 3:14, the title does not require that Christ is the first creature. Read with John 1 and Colossians 1, it means that Christ is the origin and ruler of creation: creation begins from Him because all things came to be through Him.",
            None,
        ),
        ("Proverbs 8 and Created Wisdom", "Heading 3"),
        (
            "Arian and modern anti-trinitarian arguments often appeal to Proverbs 8, where Wisdom says she was “created” or “possessed” at the beginning, depending on translation. Orthodox interpretation must be careful here. Proverbs is poetic wisdom literature, not a simple metaphysical biography of the Son. The New Testament identifies Christ as God's Wisdom, but it also teaches that He is eternal, Creator, and before all things. A poetic personification cannot overturn the apostolic witness that the Word was already God in the beginning.",
            None,
        ),
        ("“Son of God” Means Lesser or Created", "Heading 3"),
        (
            "Human sonship begins in time and depends on biological generation; divine sonship does not. When Christians confess the Son as “begotten, not made,” they are distinguishing eternal generation from creation. The Son is from the Father, but not after the Father; begotten of the Father's essence, not manufactured out of nothing. A human son shares human nature with his father; the eternal Son shares the divine nature with the Father.",
            None,
        ),
        ("“Worship” Can Mean Mere Respect", "Heading 3"),
        (
            "It is true that some biblical gestures of bowing can express human respect. But the argument for Christ's deity does not rest on one ambiguous bow. It rests on the whole pattern: Jesus receives worship after divine acts, angels are commanded to worship Him, the Lamb receives heavenly doxology with the One on the throne, Christians pray to Him, and pagan observers notice hymns to Christ as divine. Taken together, the evidence exceeds ordinary honor.",
            None,
        ),
    ],
)

# Patristic and conciliar context.
liturgical_voice = find_heading(doc, "The Liturgical Voice", 3)
block_before(
    liturgical_voice,
    [
        ("Nicaea as Guardrail, Not Invention", "Heading 3"),
        (
            "The Council of Nicaea did not invent the divinity of Christ; it guarded the apostolic confession against a new reduction. Its language “God of God, Light of Light, very God of very God, begotten, not made, being of one substance with the Father” was a fence around the worship and Scripture the Church already possessed. The council's logic was simple: if the Son is a creature, He cannot save, reveal, deify, or receive worship as the Church has always confessed.",
            None,
        ),
    ],
)

# New compact apologetic section before conclusion.
conclusion = find_heading(doc, "XV. Conclusion", 1)
block_before(
    conclusion,
    [
        ("XIV. A Compact Apologetic Case", "Heading 1"),
        (
            "For website use, the argument can be presented as a cumulative case rather than a single proof-text. The Christian claim is not that one isolated verse can be made to sound trinitarian; it is that every major stream of biblical evidence converges on the same confession.",
            None,
        ),
        ("1. Jesus Is on the Creator Side of Reality", "Heading 2"),
        (
            "John 1, 1 Corinthians 8, Colossians 1, and Hebrews 1 all place the Son in the work of creation and preservation. If all things were made through Him, then He is not one of the things made.",
            None,
        ),
        ("2. Jesus Receives the Worship Due to God", "Heading 2"),
        (
            "The New Testament rejects worship of creatures, yet Jesus receives worship from disciples, angels, and the heavenly liturgy. Revelation's worship of the Lamb is especially decisive because heaven is not confused about idolatry.",
            None,
        ),
        ("3. Jesus Bears the Divine Name and Titles", "Heading 2"),
        (
            "Jesus is called God, Lord, First and Last, I AM, Son of Man enthroned with the Ancient of Days, and the one before whom every knee bows. These are not merely compliments; they are Old Testament divine identifiers applied to Christ.",
            None,
        ),
        ("4. Jesus Performs Divine Works", "Heading 2"),
        (
            "He forgives sins, gives life, judges the world, rules the Sabbath, calms the sea, creates, sustains, and sends the Spirit. These works reveal who He is.",
            None,
        ),
        ("5. Jesus Is Prayed To and Confessed Liturgically", "Heading 2"),
        (
            "The apostolic Church calls on His name, prays Maranatha, baptizes into the triune Name, breaks the Eucharistic bread as communion in His body and blood, and sings hymns to Him.",
            None,
        ),
        ("6. The Fathers Preserve, Rather Than Replace, the Apostolic Faith", "Heading 2"),
        (
            "Ignatius, Justin, Irenaeus, Tertullian, Athanasius, Cyril, and the Nicene fathers do not create a new Jesus. They defend the Jesus already worshiped in Scripture: true God, true man, one Lord Jesus Christ, the only-begotten Son of the Father.",
            None,
        ),
    ],
)

# Source additions.
sources = find_heading(doc, "Sources Consulted", 1)
block_after(
    sources,
    [
        (
            "Additional Scripture research: Philippians 2:5–11 and Isaiah 45:22–25 on universal worship; 1 Corinthians 8:4–6 on the Father and the one Lord Jesus Christ; Colossians 1:15–20 on Christ as Creator, sustainer, and reconciler; Acts 7:59, 1 Corinthians 1:2, 1 Corinthians 16:22, and Revelation 22:20 on invocation and prayer to Christ.",
            None,
        ),
        (
            "Additional web sources checked: BibleGateway passages for Philippians 2, Isaiah 45, 1 Corinthians 8, and Colossians 1; Fordham Internet History Sourcebook, Pliny Letters 10.96–97 (https://sourcebooks.fordham.edu/source/pliny1.asp); St. Justin Martyr, First Apology (https://www.newadvent.org/fathers/0126.htm); First Council of Nicaea (https://www.newadvent.org/fathers/3801.htm).",
            None,
        ),
        (
            "Modern research orientation: Larry Hurtado's work on earliest Christian devotion to Jesus and Richard Bauckham's divine-identity framing were used as background categories for organizing the biblical material, while the document's claims are argued from Scripture and primary early sources.",
            None,
        ),
    ],
)

for table in doc.tables:
    table.autofit = True

# Clean up an empty manual page-break paragraph from the original file that
# renders as a visible square in macOS Quick Look thumbnails.
for p in list(doc.paragraphs):
    if not text(p):
        breaks = p._p.findall(".//" + qn("w:br"))
        if any(br.get(qn("w:type")) == "page" for br in breaks):
            p._p.getparent().remove(p._p)

doc.save(OUT)
print(OUT)
